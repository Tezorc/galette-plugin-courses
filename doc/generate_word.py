"""
Génère le mode d'emploi du plugin Galette Courses au format Word (.docx).
Les captures d'écran sont représentées par des encadrés gris avec légende.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_placeholder(doc, label, height_cm=5):
    """Ajoute un encadré gris clair en guise de capture d'écran."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E8E8E8')
    set_cell_borders(cell, '999999')
    cell.width = Cm(15)

    inner = cell.paragraphs[0]
    inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inner.paragraph_format.space_before = Pt(18)
    inner.paragraph_format.space_after = Pt(6)
    run = inner.add_run('📷  ' + label)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True

    # Hauteur simulée par des sauts de ligne
    lines = max(1, int(height_cm * 1.5))
    for _ in range(lines):
        p = cell.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(10)

    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(10)
    r = note.add_run('[ Insérer la capture d\'écran ]')
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()  # espace après


def apply_inline(para, text):
    """Applique le gras (**...**) et le code inline (`...`) dans un paragraphe."""
    pattern = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(1) is not None:          # gras
            run = para.add_run(m.group(1))
            run.bold = True
        else:                               # code inline
            run = para.add_run(m.group(2))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def add_paragraph_inline(doc, text, style='Normal', indent=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(indent)
    apply_inline(p, text)
    return p


def style_exists(doc, name):
    return any(s.name == name for s in doc.styles)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def create_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    # Styles de base
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    for lvl, size, bold, color in [
        ('Heading 1', 18, True,  '1F3864'),
        ('Heading 2', 14, True,  '2E74B5'),
        ('Heading 3', 11, True,  '2E74B5'),
    ]:
        s = doc.styles[lvl]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        s.paragraph_format.space_before = Pt(12 if lvl != 'Heading 1' else 20)
        s.paragraph_format.space_after  = Pt(4)

    return doc


# ---------------------------------------------------------------------------
# Table parser
# ---------------------------------------------------------------------------

def parse_md_table(lines, start):
    """Retourne (rows, next_index). rows[0] = en-têtes."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[-| :]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_md_table(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            # Nettoyer le contenu par défaut
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            apply_inline(p, cell_text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if r_idx == 0:
                set_cell_bg(cell, '2E74B5')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
            elif r_idx % 2 == 0:
                set_cell_bg(cell, 'EBF3FB')

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Image placeholders — injectés après certaines sections
# ---------------------------------------------------------------------------

PLACEHOLDERS = {
    '## Presentation': (
        'Vue d\'ensemble — tableau de bord Galette avec menus du plugin', 4),
    '### 1. Creer un evenement': (
        'Formulaire de création d\'un événement (nom, type, lieu, créneaux)', 6),
    '#### Evenement recurrent': (
        'Section "Planification" avec recurrence cochée et champs de recurrence', 5),
    '### Statuts des evenements': (
        'Page de détail d\'un événement avec boutons de workflow (Soumettre / Valider / Rejeter)', 5),
    '### 3. Consulter les seances': (
        'Liste des séances — grille de cards colorées avec jauge de remplissage', 6),
    '### 4. S\'inscrire a une seance (membre)': (
        'Page de détail d\'une séance — en-tête coloré, jauge, bouton S\'inscrire', 6),
    '### 6. Consulter ses inscriptions (membre)': (
        'Page "Mes inscriptions" — prochaine séance mise en avant, grille à venir', 5),
    '### 8. Liste d\'attente': (
        'Page de détail d\'une séance pleine — message jaune + bouton liste d\'attente', 4),
    '### 10. Statistiques (staff / admin)': (
        'Page Statistiques — cards compteurs, graphiques Chart.js, tableau taux de remplissage', 6),
    '### 12. Gestion des moniteurs': (
        'Section Moniteurs sur la page de détail d\'une séance', 4),
    '### 15. Modifier une seance (staff/admin)': (
        'Formulaire de modification de séance (date, horaire, capacité)', 4),
    '### 17. Pointage des presences (moniteur/staff)': (
        'Section pointage des présences — liste des inscrits avec select de statut', 5),
    '## Workflow de validation': (
        'Page de détail d\'un événement En attente — boutons Valider / Rejeter visibles', 5),
    '## Preferences du plugin (staff / admin)': (
        'Page Préférences — sections Dates de fermeture et Génération automatique', 5),
    '## Modeles de courriels (admin)': (
        'Page Modèles de courriels — liste des 10 modèles avec éditeur sujet/corps', 5),
    '## Navigation et menus': (
        'Barre latérale Galette — menus "Mes inscriptions" et "Gestion des inscriptions"', 4),
}


# ---------------------------------------------------------------------------
# Conversion principale
# ---------------------------------------------------------------------------

def convert(md_path, out_path):
    with open(md_path, encoding='utf-8') as f:
        raw = f.read()

    lines = raw.splitlines()
    doc   = create_doc()

    # Page de titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Cm(3)
    r = title.add_run('Plugin Galette Courses')
    r.font.name = 'Calibri'
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('Guide d\'utilisation')
    rs.font.size = Pt(16)
    rs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = ver.add_run('Version 0.1.0 — Avril 2026')
    rv.font.size = Pt(10)
    rv.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_placeholder(doc, 'Logo ou bannière de l\'application', 3)
    doc.add_page_break()

    i = 0
    in_code_block = False
    code_lines   = []
    list_depth   = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Bloc de code ---
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # Fin du bloc
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                # fond gris via shading sur le paragraphe
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F0F0')
                pPr.append(shd)
                i += 1
                continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Ligne vide ---
        if not stripped:
            i += 1
            continue

        # --- Séparateur ---
        if stripped == '---':
            doc.add_paragraph().add_run('').font.size = Pt(4)
            i += 1
            continue

        # --- Titres ---
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if m:
                lvl   = len(m.group(1))
                title_text = m.group(2)
                key   = stripped

                style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
                p = doc.add_paragraph(style=style_map.get(lvl, 'Heading 3'))
                apply_inline(p, title_text)

                # Saut de page avant H1 (sauf tout début)
                if lvl == 1 and doc.paragraphs and len(doc.paragraphs) > 10:
                    p.paragraph_format.page_break_before = True

                # Placeholder après certains titres
                for ph_key, (ph_label, ph_h) in PLACEHOLDERS.items():
                    if key == ph_key:
                        add_placeholder(doc, ph_label, ph_h)
                        break

                i += 1
                continue

        # --- Tableau Markdown ---
        if stripped.startswith('|'):
            rows, i = parse_md_table(lines, i)
            add_md_table(doc, rows)
            continue

        # --- Liste numérotée ---
        m_num = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m_num:
            indent_spaces = len(m_num.group(1))
            indent_cm = 0.5 + indent_spaces * 0.15
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent   = Cm(indent_cm)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(1)
            apply_inline(p, m_num.group(3))
            i += 1
            continue

        # --- Liste à puces ---
        m_bul = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m_bul:
            indent_spaces = len(m_bul.group(1))
            indent_cm = 0.3 + indent_spaces * 0.3
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Cm(indent_cm)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            apply_inline(p, m_bul.group(2))
            i += 1
            continue

        # --- Paragraphe normal ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        apply_inline(p, stripped)
        i += 1

    # Pied de page
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('Plugin Galette Courses — Guide d\'utilisation — ccag42 Team')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(out_path)
    print(f'Document genere : {out_path}')


# ---------------------------------------------------------------------------
if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    convert(
        os.path.join(base, 'mode-emploi.md'),
        os.path.join(base, 'mode-emploi.docx'),
    )
